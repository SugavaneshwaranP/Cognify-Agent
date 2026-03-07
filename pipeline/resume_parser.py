"""
CognifyX – Resume Parser
Extracts text from PDF, DOCX, and TXT files with cleanup.
Supports both directory parsing and in-memory file handling.
"""
import pypdf
import os
import io
import re
import docx


class ResumeParser:
    @staticmethod
    def clean_text(text):
        """Fixes issues with spaced-out text and preserves line structure."""
        if not text:
            return ""

        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line_strip = line.strip()
            if not line_strip:
                continue

            # Detect spaced-out text (e.g., "S K I L L S")
            if len(re.findall(r'[a-zA-Z]\s', line_strip)) > len(line_strip) / 3:
                words = re.split(r'\s{2,}', line_strip)
                cleaned_words = [re.sub(r'\s', '', w) for w in words]
                line_strip = " ".join(cleaned_words)

            cleaned_lines.append(line_strip)

        return "\n".join(cleaned_lines)

    @staticmethod
    def extract_text(file_path):
        """Extracts text from PDF, DOCX, or TXT files."""
        ext = os.path.splitext(file_path)[1].lower()

        try:
            text = ""
            if ext == '.pdf':
                with open(file_path, 'rb') as file:
                    reader = pypdf.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text() or ""
            elif ext == '.docx':
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()

            return ResumeParser.clean_text(text)
        except Exception as e:
            print(f"Error parsing {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_bytes(file_bytes, filename):
        """Extract text from in-memory file bytes (useful for Streamlit uploads)."""
        ext = os.path.splitext(filename)[1].lower()
        try:
            text = ""
            if ext == '.pdf':
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    text += page.extract_text() or ""
            elif ext == '.docx':
                doc = docx.Document(io.BytesIO(file_bytes))
                text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == '.txt':
                text = file_bytes.decode('utf-8', errors='ignore')

            return ResumeParser.clean_text(text)
        except Exception as e:
            print(f"Error parsing {filename}: {e}")
            return ""

    @staticmethod
    def parse_directory(directory_path):
        """Parses all resumes in a directory (recursive)."""
        resumes = []
        valid_extensions = ('.pdf', '.docx', '.txt')

        for root, _, files in os.walk(directory_path):
            for file in files:
                if file.lower().endswith(valid_extensions):
                    path = os.path.join(root, file)
                    text = ResumeParser.extract_text(path)
                    if text and len(text.strip()) > 50:  # Skip near-empty files
                        resumes.append({
                            "filename": file,
                            "path": path,
                            "text": text
                        })
        return resumes
