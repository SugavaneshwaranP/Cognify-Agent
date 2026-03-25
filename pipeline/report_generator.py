"""
CognifyX – Report Generator
Creates professional PDF and DOCX hiring reports with charts and tables.
Uses reportlab for PDF, python-docx for DOCX, matplotlib for charts.
"""
import io
import os
import tempfile
from datetime import datetime

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image as RLImage, PageBreak, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.ticker as mticker
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False


# ── Colour palette ──────────────────────────────────────────────────
BRAND_DARK   = "#1e1b4b"
BRAND_MID    = "#4338ca"
BRAND_LIGHT  = "#a5b4fc"
ACCENT_CYAN  = "#06b6d4"
ACCENT_GREEN = "#10b981"
ACCENT_AMBER = "#f59e0b"
ACCENT_RED   = "#ef4444"
TEXT_DARK    = "#1e293b"
TEXT_MID     = "#475569"
TEXT_LIGHT   = "#94a3b8"


def _hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple."""
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


# ════════════════════════════════════════════════════════════════════
# CHART GENERATORS (matplotlib)
# ════════════════════════════════════════════════════════════════════

def _create_score_comparison_chart(matrix_data):
    """Bar chart comparing candidate scores."""
    if not HAS_MATPLOTLIB or not matrix_data:
        return None

    rows = matrix_data.get('rows', [])
    if not rows:
        return None

    names = [r.get('filename', 'Unknown')[:15] for r in rows]
    ats = [r.get('ats_score', 0) for r in rows]
    kw =  [r.get('keyword_score', 0) for r in rows]
    llm = [r.get('llm_score', 0) for r in rows]
    comp = [r.get('composite_score', 0) for r in rows]

    fig, ax = plt.subplots(figsize=(8, 4))
    x = range(len(names))
    w = 0.20

    ax.bar([i - 1.5*w for i in x], ats, w, label='ATS Score', color=ACCENT_CYAN, alpha=0.85)
    ax.bar([i - 0.5*w for i in x], kw,  w, label='Keyword %',  color=BRAND_MID, alpha=0.85)
    ax.bar([i + 0.5*w for i in x], llm, w, label='LLM Score',  color=ACCENT_AMBER, alpha=0.85)
    ax.bar([i + 1.5*w for i in x], comp, w, label='Composite', color=ACCENT_GREEN, alpha=0.85)

    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=25, ha='right', fontsize=8)
    ax.set_ylabel('Score')
    ax.set_title('Candidate Score Comparison', fontweight='bold', fontsize=11)
    ax.legend(loc='upper right', fontsize=7)
    ax.set_ylim(0, 110)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _create_skill_gap_chart(skill_gap_data):
    """Horizontal bar chart for skill coverage."""
    if not HAS_MATPLOTLIB or not skill_gap_data:
        return None

    coverage = skill_gap_data.get('coverage', [])[:12]
    if not coverage:
        return None

    keywords = [c['keyword'][:18] for c in coverage]
    pcts = [c['coverage_pct'] for c in coverage]
    bar_colors = [
        ACCENT_GREEN if p >= 70 else ACCENT_AMBER if p >= 40 else ACCENT_RED
        for p in pcts
    ]

    fig, ax = plt.subplots(figsize=(7, max(3, len(keywords) * 0.35)))
    bars = ax.barh(keywords, pcts, color=bar_colors, alpha=0.85, height=0.6)
    ax.set_xlim(0, 110)
    ax.set_xlabel('Coverage %')
    ax.set_title('Skill Coverage Across Candidate Pool', fontweight='bold', fontsize=11)
    ax.axvline(x=70, color=TEXT_LIGHT, linestyle='--', linewidth=0.8, alpha=0.6)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    for bar, pct in zip(bars, pcts):
        ax.text(bar.get_width() + 2, bar.get_y() + bar.get_height()/2,
                f'{pct}%', va='center', fontsize=8, color=TEXT_MID)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _create_diversity_chart(diversity_data):
    """Pie charts for diversity metrics."""
    if not HAS_MATPLOTLIB or not diversity_data:
        return None

    fig, axes = plt.subplots(1, 3, figsize=(10, 3.5))
    pie_colors = [ACCENT_CYAN, BRAND_MID, ACCENT_AMBER, ACCENT_GREEN, ACCENT_RED, TEXT_LIGHT]

    # Domain distribution
    domains = diversity_data.get('domain_distribution', {})
    if domains:
        labels = list(domains.keys())[:6]
        sizes = list(domains.values())[:6]
        axes[0].pie(sizes, labels=labels, autopct='%1.0f%%', startangle=90,
                    colors=pie_colors[:len(labels)], textprops={'fontsize': 7})
        axes[0].set_title('Domains', fontsize=9, fontweight='bold')

    # Education
    edu = diversity_data.get('education_distribution', {})
    if edu:
        labels = [k for k, v in edu.items() if v > 0]
        sizes = [v for v in edu.values() if v > 0]
        if sizes:
            axes[1].pie(sizes, labels=labels, autopct='%1.0f%%', startangle=90,
                        colors=pie_colors[:len(labels)], textprops={'fontsize': 7})
        axes[1].set_title('Education', fontsize=9, fontweight='bold')

    # Experience
    exp = diversity_data.get('experience_distribution', {})
    if exp:
        labels = [k for k, v in exp.items() if v > 0]
        sizes = [v for v in exp.values() if v > 0]
        if sizes:
            axes[2].pie(sizes, labels=labels, autopct='%1.0f%%', startangle=90,
                        colors=pie_colors[:len(labels)], textprops={'fontsize': 7})
        axes[2].set_title('Experience', fontsize=9, fontweight='bold')

    plt.suptitle('Diversity & Inclusion Metrics', fontsize=11, fontweight='bold', y=1.02)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _create_salary_chart(salary_data):
    """Range chart for salary benchmarks."""
    if not HAS_MATPLOTLIB or not salary_data:
        return None

    benchmarks = salary_data.get('benchmarks', [])
    if not benchmarks:
        return None

    names = [b.get('filename', 'Unknown')[:15] for b in benchmarks]
    mins = [int(b['salary_min'].replace('$', '').replace(',', '')) for b in benchmarks]
    maxs = [int(b['salary_max'].replace('$', '').replace(',', '')) for b in benchmarks]
    mids = [int(b['midpoint'].replace('$', '').replace(',', '')) for b in benchmarks]

    fig, ax = plt.subplots(figsize=(7, max(3, len(names) * 0.4)))
    y = range(len(names))

    for i in y:
        ax.barh(i, maxs[i] - mins[i], left=mins[i], height=0.5,
                color=ACCENT_CYAN, alpha=0.3)
        ax.plot(mids[i], i, 'o', color=BRAND_MID, markersize=8)

    ax.set_yticks(y)
    ax.set_yticklabels(names, fontsize=8)
    ax.set_xlabel('Annual Salary (USD)')
    ax.set_title('Salary Range Benchmarks', fontweight='bold', fontsize=11)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f'${x:,.0f}'))
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════════════════
# PDF GENERATOR
# ════════════════════════════════════════════════════════════════════

class PDFReportGenerator:
    """Generates a professional PDF hiring report using reportlab."""

    def __init__(self):
        if not HAS_REPORTLAB:
            raise ImportError("reportlab is required: pip install reportlab")

    def generate(self, report_data):
        """Create a PDF report and return it as bytes."""
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                topMargin=20*mm, bottomMargin=20*mm,
                                leftMargin=18*mm, rightMargin=18*mm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('Title', parent=styles['Title'],
                                      fontSize=22, spaceAfter=6,
                                      textColor=colors.HexColor(BRAND_DARK))
        heading_style = ParagraphStyle('H2', parent=styles['Heading2'],
                                        fontSize=14, spaceBefore=18, spaceAfter=8,
                                        textColor=colors.HexColor(BRAND_MID))
        body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                     fontSize=9.5, leading=13,
                                     textColor=colors.HexColor(TEXT_DARK))
        small_style = ParagraphStyle('Small', parent=styles['Normal'],
                                      fontSize=8, textColor=colors.HexColor(TEXT_MID))

        elements = []

        # ── Cover ──
        elements.append(Spacer(1, 30*mm))
        elements.append(Paragraph("🧠 CognifyX", title_style))
        elements.append(Paragraph("AI Hiring Intelligence Report", heading_style))
        gen_at = report_data.get('generated_at', datetime.now().strftime("%Y-%m-%d %H:%M"))
        elements.append(Paragraph(f"Generated: {gen_at}", small_style))
        elements.append(Spacer(1, 10*mm))
        elements.append(HRFlowable(width="100%", thickness=1,
                                    color=colors.HexColor(BRAND_LIGHT)))
        elements.append(Spacer(1, 8*mm))

        # ── Executive Summary ──
        exec_sum = report_data.get('executive_summary', {})
        if exec_sum:
            elements.append(Paragraph("📋 Executive Summary", heading_style))
            elements.append(Paragraph(exec_sum.get('content', ''), body_style))
            stats = exec_sum.get('stats', {})
            stats_text = (
                f"Screened: {stats.get('total_screened', 0)} | "
                f"Filtered: {stats.get('ats_filtered', 0)} | "
                f"Shortlisted: {stats.get('shortlisted', 0)} | "
                f"Top: {stats.get('top_candidate', 'N/A')} ({stats.get('top_score', 0)}/100)"
            )
            elements.append(Paragraph(stats_text, small_style))
            elements.append(Spacer(1, 6*mm))

        # ── Comparison Matrix ──
        matrix = report_data.get('comparison_matrix', {})
        rows = matrix.get('rows', [])
        if rows:
            elements.append(Paragraph("📊 Candidate Comparison Matrix", heading_style))

            table_data = [["#", "File", "Domain", "Exp", "ATS", "KW%",
                           "LLM", "Composite"]]
            for r in rows:
                table_data.append([
                    str(r['rank']), r.get('filename', 'Unknown')[:20], r['domain'][:15],
                    str(r['experience_years']), str(r['ats_score']),
                    str(r['keyword_score']), str(r['llm_score']),
                    str(r['composite_score']),
                ])

            t = Table(table_data, repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BRAND_DARK)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('FONTSIZE', (0, 0), (-1, 0), 8.5),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(TEXT_LIGHT)),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                 [colors.white, colors.HexColor('#f8fafc')]),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 6*mm))

            # Score chart
            chart_buf = _create_score_comparison_chart(matrix)
            if chart_buf:
                elements.append(RLImage(chart_buf, width=6.5*inch, height=3*inch))
                elements.append(Spacer(1, 6*mm))

        # ── Skill Gap Analysis ──
        skill_gap = report_data.get('skill_gap_analysis', {})
        if skill_gap:
            elements.append(PageBreak())
            elements.append(Paragraph("🔍 Skill Gap Analysis", heading_style))

            coverage = skill_gap.get('coverage', [])
            if coverage:
                gap_table = [["Keyword", "Candidates", "Coverage", "Status"]]
                for c in coverage[:15]:
                    gap_table.append([
                        c['keyword'], str(c['candidates_with']),
                        f"{c['coverage_pct']}%", c['status']
                    ])
                gt = Table(gap_table, repeatRows=1)
                gt.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BRAND_DARK)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(TEXT_LIGHT)),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                     [colors.white, colors.HexColor('#f8fafc')]),
                ]))
                elements.append(gt)
                elements.append(Spacer(1, 4*mm))

            chart_buf = _create_skill_gap_chart(skill_gap)
            if chart_buf:
                elements.append(RLImage(chart_buf, width=5.5*inch, height=3*inch))
                elements.append(Spacer(1, 6*mm))

        # ── Interview Questions ──
        iq = report_data.get('interview_questions', {})
        if iq:
            elements.append(PageBreak())
            elements.append(Paragraph("💬 Interview Question Bank", heading_style))

            for cand in iq.get('candidates', []):
                elements.append(Paragraph(
                    f"<b>#{cand['rank']} – {cand.get('candidate', 'Unknown')}</b>", body_style))
                for q in cand['questions']:
                    elements.append(Paragraph(
                        f"  • [{q['type']}] {q['question']}", small_style))
                    elements.append(Paragraph(
                        f"    <i>Reason: {q['reason']}</i>", small_style))
                elements.append(Spacer(1, 3*mm))

        # ── Diversity ──
        diversity = report_data.get('diversity_metrics', {})
        if diversity:
            elements.append(PageBreak())
            elements.append(Paragraph("🌍 Diversity & Inclusion Metrics", heading_style))

            di = diversity.get('diversity_index', 0)
            elements.append(Paragraph(
                f"Diversity Index: <b>{di}%</b> | "
                f"Unique Domains: {diversity.get('unique_domains', 0)} | "
                f"Total Pool: {diversity.get('total_pool', 0)}",
                body_style
            ))
            elements.append(Spacer(1, 3*mm))

            chart_buf = _create_diversity_chart(diversity)
            if chart_buf:
                elements.append(RLImage(chart_buf, width=6.5*inch, height=2.8*inch))

        # ── Salary Benchmarks ──
        salary = report_data.get('salary_benchmarks', {})
        benchmarks = salary.get('benchmarks', [])
        if benchmarks:
            elements.append(PageBreak())
            elements.append(Paragraph("💰 Salary Benchmarking", heading_style))

            sal_table = [["#", "File", "Level", "Min", "Max", "Midpoint"]]
            for b in benchmarks:
                sal_table.append([
                    str(b['rank']), b.get('filename', 'Unknown')[:20], b['level'],
                    b['salary_min'], b['salary_max'], b['midpoint']
                ])
            st = Table(sal_table, repeatRows=1)
            st.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BRAND_DARK)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(TEXT_LIGHT)),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                 [colors.white, colors.HexColor('#f8fafc')]),
            ]))
            elements.append(st)
            elements.append(Spacer(1, 4*mm))

            chart_buf = _create_salary_chart(salary)
            if chart_buf:
                elements.append(RLImage(chart_buf, width=5.5*inch, height=3*inch))

        elements.append(Spacer(1, 10*mm))
        elements.append(Paragraph(salary.get('note', ''), small_style))

        # ── Footer ──
        elements.append(Spacer(1, 10*mm))
        elements.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor(TEXT_LIGHT)))
        elements.append(Paragraph(
            f"© {datetime.now().year} CognifyX | Generated {gen_at}",
            ParagraphStyle('Footer', parent=small_style, alignment=TA_CENTER, fontSize=7)
        ))

        doc.build(elements)
        buf.seek(0)
        return buf.getvalue()


# ════════════════════════════════════════════════════════════════════
# DOCX GENERATOR
# ════════════════════════════════════════════════════════════════════

class DOCXReportGenerator:
    """Generates a professional DOCX hiring report using python-docx."""

    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx is required: pip install python-docx")

    def _add_heading(self, doc, text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*_hex_to_rgb(BRAND_DARK))
        return h

    def _add_table(self, doc, headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Shading Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8.5)

        return table

    def generate(self, report_data):
        """Create a DOCX report and return it as bytes."""
        doc = Document()

        # Title
        title = doc.add_heading('🧠 CognifyX – AI Hiring Intelligence Report', 0)
        for run in title.runs:
            run.font.color.rgb = RGBColor(*_hex_to_rgb(BRAND_DARK))

        gen_at = report_data.get('generated_at', datetime.now().strftime("%Y-%m-%d %H:%M"))
        p = doc.add_paragraph(f'Generated: {gen_at}')
        p.style = doc.styles['Normal']
        doc.add_paragraph('─' * 60)

        # ── Executive Summary ──
        exec_sum = report_data.get('executive_summary', {})
        if exec_sum:
            self._add_heading(doc, '📋 Executive Summary', level=1)
            doc.add_paragraph(exec_sum.get('content', ''))
            stats = exec_sum.get('stats', {})
            doc.add_paragraph(
                f"Screened: {stats.get('total_screened', 0)} | "
                f"Filtered: {stats.get('ats_filtered', 0)} | "
                f"Shortlisted: {stats.get('shortlisted', 0)}"
            )

        # ── Comparison Matrix ──
        matrix = report_data.get('comparison_matrix', {})
        rows = matrix.get('rows', [])
        if rows:
            self._add_heading(doc, '📊 Candidate Comparison Matrix', level=1)
            headers = ["#", "File", "Domain", "Exp", "ATS", "KW%", "LLM", "Composite"]
            table_rows = []
            for r in rows:
                table_rows.append([
                    str(r['rank']), r.get('filename', 'Unknown')[:25], r['domain'][:15],
                    str(r['experience_years']), str(r['ats_score']),
                    str(r['keyword_score']), str(r['llm_score']),
                    str(r['composite_score']),
                ])
            self._add_table(doc, headers, table_rows)

            # Add chart
            chart_buf = _create_score_comparison_chart(matrix)
            if chart_buf:
                doc.add_paragraph()
                doc.add_picture(chart_buf, width=Inches(6))

        # ── Skill Gap ──
        skill_gap = report_data.get('skill_gap_analysis', {})
        if skill_gap:
            doc.add_page_break()
            self._add_heading(doc, '🔍 Skill Gap Analysis', level=1)
            coverage = skill_gap.get('coverage', [])
            if coverage:
                self._add_table(doc,
                    ["Keyword", "Candidates", "Coverage", "Status"],
                    [[c['keyword'], str(c['candidates_with']),
                      f"{c['coverage_pct']}%", c['status']] for c in coverage[:15]]
                )
            chart_buf = _create_skill_gap_chart(skill_gap)
            if chart_buf:
                doc.add_paragraph()
                doc.add_picture(chart_buf, width=Inches(5.5))

        # ── Interview Questions ──
        iq = report_data.get('interview_questions', {})
        if iq:
            doc.add_page_break()
            self._add_heading(doc, '💬 Interview Question Bank', level=1)
            for cand in iq.get('candidates', []):
                self._add_heading(doc, f"#{cand['rank']} – {cand.get('candidate', 'Unknown')}", level=2)
                for q in cand['questions']:
                    doc.add_paragraph(
                        f"[{q['type']}] {q['question']}",
                        style='List Bullet'
                    )
                    reason_p = doc.add_paragraph(f"Reason: {q['reason']}")
                    reason_p.style = doc.styles['Normal']
                    for run in reason_p.runs:
                        run.font.italic = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(*_hex_to_rgb(TEXT_MID))

        # ── Diversity ──
        diversity = report_data.get('diversity_metrics', {})
        if diversity:
            doc.add_page_break()
            self._add_heading(doc, '🌍 Diversity & Inclusion Metrics', level=1)
            doc.add_paragraph(
                f"Diversity Index: {diversity.get('diversity_index', 0)}% | "
                f"Unique Domains: {diversity.get('unique_domains', 0)} | "
                f"Total Pool: {diversity.get('total_pool', 0)}"
            )
            chart_buf = _create_diversity_chart(diversity)
            if chart_buf:
                doc.add_picture(chart_buf, width=Inches(6))

        # ── Salary ──
        salary = report_data.get('salary_benchmarks', {})
        benchmarks = salary.get('benchmarks', [])
        if benchmarks:
            doc.add_page_break()
            self._add_heading(doc, '💰 Salary Benchmarking', level=1)
            self._add_table(doc,
                ["#", "File", "Level", "Min", "Max", "Midpoint"],
                [[str(b['rank']), b.get('filename', 'Unknown')[:25], b['level'],
                  b['salary_min'], b['salary_max'], b['midpoint']]
                 for b in benchmarks]
            )
            chart_buf = _create_salary_chart(salary)
            if chart_buf:
                doc.add_paragraph()
                doc.add_picture(chart_buf, width=Inches(5.5))
            doc.add_paragraph(salary.get('note', ''))

        # ── Footer ──
        doc.add_paragraph('─' * 60)
        footer = doc.add_paragraph(
            f"© {datetime.now().year} CognifyX | Generated {gen_at}"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(*_hex_to_rgb(TEXT_LIGHT))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
